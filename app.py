from flask import Flask, render_template, request, send_from_directory
import os
import markdown
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import google.generativeai as genai
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load API key from .env file
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Flask setup
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs("generated_reports", exist_ok=True)

def clean_markdown_to_text(markdown_text):
    lines = markdown_text.split('\n')
    clean_lines = []
    for line in lines:
        if line.startswith('#'):
            line = line.lstrip('#').strip()
        if line.startswith('* ') or line.startswith('- '):
            line = "• " + line[2:]
        line = line.replace('**', '')
        clean_lines.append(line)
    return clean_lines

def generate_report_file(markdown_text, filename="cyber_threat_report", format="pdf"):
    lines = clean_markdown_to_text(markdown_text)
    file_path = f"generated_reports/{filename}.{format}"

    if format == "pdf":
        class PDF(FPDF):
            def header(self):
                self.set_font("Times", 'I', 8)
                self.cell(0, 10, "", 0, 1, 'C')

            def footer(self):
                self.set_y(-15)
                self.set_font('Times', 'I', 8)
                self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

        pdf = PDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_left_margin(20)
        pdf.set_right_margin(20)
        pdf.set_font("Times", size=11)

        # Title
        pdf.set_font("Times", 'B', 16)
        pdf.cell(0, 10, "Cyber Threat Intelligence Report: Global Cybersecurity Incidents (2015-2024)", ln=True, align="C")
        pdf.ln(5)
        pdf.set_font("Times", size=11)

        for line in lines:
            line = line.strip()
            if not line:
                pdf.ln(5)
            elif line.endswith(':') and not line.startswith("•"):
                pdf.set_font("Times", 'B', 14)
                pdf.multi_cell(0, 10, txt=line)
                pdf.set_font("Times", size=11)
            else:
                pdf.multi_cell(0, 10, txt=line)

        pdf.output(file_path)

    elif format == "docx":
        doc = Document()

        # Set margins (in points: 1 inch = 72 pt)
        for section in doc.sections:
            section.top_margin = Pt(36)
            section.bottom_margin = Pt(36)
            section.left_margin = Pt(48)
            section.right_margin = Pt(48)

        # Title
        title = doc.add_paragraph()
        run = title.add_run("Cyber Threat Intelligence Report: Global Cybersecurity Incidents (2015-2024)")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacer

        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.endswith(':') and not line.startswith("•"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = "Times New Roman"
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

        doc.save(file_path)

    return file_path

@app.route('/generate-report', methods=['POST'])
def generate_report():
    return upload()


@app.route('/')
def index():
    return render_template('index.html')

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory("generated_reports", filename, as_attachment=True)

@app.route('/upload', methods=['POST'])
def upload():
    if 'logfile' not in request.files:
        return 'Error: No file part in the request.'

    file = request.files['logfile']
    if file.filename == '':
        return 'Error: No file selected for upload.'

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    with open(filepath, 'r') as f:
        log_data = f.read()

    # Gemini Prompt
    prompt = f"""
You are a cybersecurity analyst. Given the following log data, generate a detailed and professional Cyber Threat Intelligence Report of 8 to 10 pages in markdown covering:
1. Executive Summary,
2. Threat Overview,
3. Threat Intelligence Findings,
4. Data Sources & Collection,
5. Victimology,
6. Impact Assessment,
7. Attack Lifecycle (Kill Chain or MITRE ATT&CK Mapping),
8. Analysis & Attribution,
9. Mitigation & Recommendations,
10. Incident Response Guidance,
11. Appendices & References.

Log Data:
{log_data}
"""

    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content(prompt)

    markdown_text = response.text
    html_report = markdown.markdown(markdown_text)

    export_format = request.form.get("format", "docx").lower()
    if export_format not in ["pdf", "docx"]:
        export_format = "docx"

    report_path = generate_report_file(markdown_text, format=export_format)
    download_filename = os.path.basename(report_path)

    return render_template('report.html', report=html_report, download_filename=download_filename)

if __name__ == '__main__':
    app.run(debug=True)
